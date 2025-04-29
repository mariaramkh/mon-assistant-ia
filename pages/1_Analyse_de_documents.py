import streamlit as st
import tempfile, os
from docx import Document
from langchain_community.document_loaders import (
    PyPDFLoader, TextLoader, Docx2txtLoader, UnstructuredExcelLoader,
    CSVLoader, JSONLoader, UnstructuredMarkdownLoader
)
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.chains import ConversationalRetrievalChain
from langchain_community.llms import Ollama
from langchain.memory import ConversationBufferMemory
from dotenv import load_dotenv
from langchain.chat_models import ChatOpenAI

load_dotenv()

st.markdown("## Analyse de documents")

uploaded_files = st.file_uploader("Importez vos fichiers", type=["pdf", "txt", "docx", "xlsx", "csv", "json", "md"], accept_multiple_files=True)
selected_model = st.selectbox("Modèle IA", ["llama3", "mistral", "gemma", "gpt-4"])

def smart_loader(path, ext):
    if ext == ".pdf": return PyPDFLoader(path)
    elif ext == ".txt": return TextLoader(path, encoding="utf-8")
    elif ext == ".docx": return Docx2txtLoader(path)
    elif ext == ".xlsx": return UnstructuredExcelLoader(path)
    elif ext == ".csv": return CSVLoader(file_path=path)
    elif ext == ".json": return JSONLoader(file_path=path)
    elif ext == ".md": return UnstructuredMarkdownLoader(path)
    else: raise ValueError("Format non supporté")

def export_to_word(messages):
    doc = Document()
    doc.add_heading("Historique de conversation IA", 0)
    for msg in messages:
        role = "Utilisateur" if msg["role"] == "user" else "Assistant"
        doc.add_paragraph(f"{role} :", style="Heading 2")
        doc.add_paragraph(msg["content"])
    path = os.path.join(tempfile.gettempdir(), "conversation.docx")
    doc.save(path)
    return path

if uploaded_files:
    all_docs = []
    for file in uploaded_files:
        ext = os.path.splitext(file.name)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(file.read())
            tmp_path = tmp.name
        try:
            loader = smart_loader(tmp_path, ext)
            documents = loader.load()
            all_docs.extend(documents)
        except Exception as e:
            st.error(f"Erreur avec {file.name} : {e}")

    if all_docs:
        splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        docs = splitter.split_documents(all_docs)
        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        db = FAISS.from_documents(docs, embeddings)

        try:
            if selected_model == "gpt-4":
                llm = ChatOpenAI(model_name="gpt-4", temperature=0)
            else:
                llm = Ollama(model=selected_model, base_url="http://localhost:11434")

            memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
            qa = ConversationalRetrievalChain.from_llm(llm=llm, retriever=db.as_retriever(), memory=memory)

            if "messages" not in st.session_state:
                st.session_state.messages = []

            for msg in st.session_state.messages:
                with st.chat_message(msg["role"]):
                    st.markdown(msg["content"])

            if prompt := st.chat_input("Pose ta question sur les documents :"):
                st.session_state.messages.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.markdown(prompt)
                with st.chat_message("assistant"):
                    with st.spinner("L'IA réfléchit..."):
                        response = qa.run({"question": "Réponds en français : " + prompt})
                        st.markdown(response)
                st.session_state.messages.append({"role": "assistant", "content": response})

            if st.download_button("Télécharger la conversation",
                data=open(export_to_word(st.session_state.messages), "rb").read(),
                file_name="conversation_IA.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
                st.success("Conversation enregistrée.")
        except Exception as e:
            st.error(f"Erreur avec le modèle sélectionné : {e}")
else:
    st.info("Veuillez importer au moins un document.")
